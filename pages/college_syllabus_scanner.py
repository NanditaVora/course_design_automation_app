import datetime
import os
from io import BytesIO
from docx import Document
import streamlit as st
from dotenv import load_dotenv

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import AzureChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
import chromadb

# Initialize session state
def initalize_session_state():
    st.session_state.response = ""
    st.session_state.content_written = False
    st.session_state.file_downloaded = False
    st.session_state.buffer = BytesIO()

# Initialize embeddings and vector store
def initialize_vector_store(uploaded_file):
    loader = PyPDFLoader(f"inputs/college-syllabi/{uploaded_file}")
    documents = loader.load()
    text = RecursiveCharacterTextSplitter().split_documents(documents)

    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5",
                                       encode_kwargs={"normalize_embeddings": True})
    chromadb.api.client.SharedSystemClient.clear_system_cache()
    vectorstore = Chroma.from_documents(documents, embedding=embeddings)
    return vectorstore

# Create the LLM chain
def create_llm_chain(vectorstore, query):
    retriever = vectorstore.as_retriever()
    load_dotenv()
    llm = AzureChatOpenAI(
        azure_deployment=os.environ.get('DEPLOYMENT_NAME_POC'),
        temperature=0,
        request_timeout=None,
        max_retries=2,
        api_key=os.environ.get('AZURE_OPENAI_API_KEY_POC'),
        api_version=os.environ.get('API_VERSION_POC'),
        azure_endpoint=os.environ.get('AZURE_OPENAI_ENDPOINT_POC')
    )


    template = """
    You are an assistant for question-answering tasks.
    Use the provided context only to answer the following question:

    <context>
    {context}
    </context>

    Question: {input}
    """
    prompt = ChatPromptTemplate.from_template(template)
    doc_chain = create_stuff_documents_chain(llm, prompt)
    return create_retrieval_chain(retriever, doc_chain)

# Process the file and query to get the response
def process_query(uploaded_file, scan_topic):
    vectorstore = initialize_vector_store(uploaded_file)
    query = f"""
        List down all the topics related to {scan_topic}.
        Ensure that you include any subtopics or specific areas covered under these main topics.
        Output Format:
        •	Main Topic 1
            o	Subtopic 1.1
            o	Subtopic 1.2
        •	Main Topic 2
            o	Subtopic 2.1
            o	Subtopic 2.2
        Example:
        •	Database Management Systems
            o	Introduction to DBMS
            o	Data Models
            o	SQL Queries
        •	Advanced Database Topics
            o	Normalization
            o	Transactions
            o	Indexing
        Please ensure the list is comprehensive and covers all relevant topics mentioned in the syllabus.

        Do not include any introductory lines or closing lines in your response, so that the response can as is downloaded.
        Thank you!
    """
    chain = create_llm_chain(vectorstore, query)
    response = chain.invoke({"input": query})
    return response['answer']

# Utility function to write or append content to a .docx file
def write_or_append_docx(file_path, from_file, content):

    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()
    doc.add_heading(f'{from_file}', 0)
    doc.add_paragraph(content)
    doc.save(file_path)
    st.session_state.content_written = True
    return doc

# Handle the Get Response button
def handle_get_response(uploaded_file, scan_topic):

    # Get Response Button
    if st.button('Get Response'):
        if uploaded_file and scan_topic:
            response = process_query(uploaded_file, scan_topic)
            st.session_state.response = response
            return response
        else:
            st.warning("Please upload a file and enter a prompt.")
            return None

# Save document in memory
def save_doc(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.session_state.buffer = buffer

# Confirm file download
def confirm_file_download():
    st.session_state.buffer = BytesIO()
    st.session_state.file_downloaded = True

# Function to get list of PDF files from a given directory
def get_pdf_files(directory):
    return [f for f in os.listdir(directory) if f.endswith('.pdf')]

# Function to create Streamlit app
def create_pdf_selector_app(pdf_directory,inputbox):
    inputbox.title("PDF File and Topic Selector")

    # Get list of PDF files
    pdf_files = get_pdf_files(pdf_directory)

    # Dropdown to select a PDF file
    selected_pdf = inputbox.selectbox("Select College Syllabus file to be Scanned", pdf_files)
    return selected_pdf

# Streamlit UI
def main():
    st.title("Welcome to Course Design Automation App")
    "**In this stage, GenAI plays the role of File Scanner and helps list topics from the selected college syllabus file.**"

    with st.sidebar:
        inputbox = st.container(height=300)
        # Directory containing PDF files
        pdf_directory = "inputs/college-syllabi"

        # Create the Streamlit app
        uploaded_file = create_pdf_selector_app(pdf_directory, inputbox)

        # scan_topic = inputbox.text_area("Enter the comma-separated keywords that will be used for scanning")
        scan_topic = inputbox.selectbox("Select topics that need to be searched in the selected college syllabus file",
                                        ["Java", "Database", "Object Oriented Programming", "OOPs", "JavaScript", "React",
                                         "Angular", "Spring"])

        handle_get_response(uploaded_file, scan_topic)
    if "response" not in st.session_state:
        initalize_session_state()

    if st.session_state.response:
        st.header("Response")
        # Save Response Button
        if st.button('Save Response to DOCX'):
            doc = write_or_append_docx(f"outputs/{scan_topic}_college_syllabus.docx", uploaded_file, st.session_state.response)
            save_doc(doc)

        if st.session_state.content_written:
            st.markdown("Content Written Successfully")
        st.markdown(st.session_state.response)







main()
